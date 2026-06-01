from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUTPUT_PATH = r"d:\WorkSpace\Sorriso Amigo\docs\plano-de-testes-sorriso-amigo.docx"


def add_title(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_case(doc, case_id, title, case_type, pre_conditions, steps, expected, priority="Alta"):
    p = doc.add_paragraph(f"{case_id} - {title}")
    p.runs[0].bold = True
    doc.add_paragraph(f"Tipo: {case_type}")
    doc.add_paragraph(f"Prioridade: {priority}")
    doc.add_paragraph("Pre-condicoes:")
    add_bullets(doc, pre_conditions)
    doc.add_paragraph("Passos:")
    add_numbers(doc, steps)
    doc.add_paragraph("Resultado esperado:")
    add_bullets(doc, expected)


def build_document():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_title(doc, "Plano de Testes Completo - Sorriso Amigo")
    doc.add_paragraph("")

    add_heading(doc, "1. Identificacao do Documento", 1)
    add_bullets(
        doc,
        [
            "Projeto: Sorriso Amigo",
            "Versao do plano: 2.0 (completo)",
            "Data: 06/04/2026",
            "Responsavel: Grupo do projeto",
            "Ambiente alvo: Producao (Render) e homologacao local"
        ],
    )

    add_heading(doc, "2. Objetivo", 1)
    doc.add_paragraph(
        "Definir a estrategia, escopo e execucao dos testes para todos os casos de uso "
        "do projeto Sorriso Amigo, assegurando qualidade funcional, consistencia de dados, "
        "seguranca de acesso e comportamento esperado em ambiente web real."
    )

    add_heading(doc, "3. Escopo de Teste", 1)
    add_bullets(
        doc,
        [
            "UC-01 - Cadastrar usuario",
            "UC-02 - Realizar login",
            "UC-03 - Registrar checklist diario",
            "UC-04 - Consultar dashboard de evolucao",
            "UC-05 - Consultar guia de escovacao",
            "UC-06 - Responder quiz educativo",
            "UC-07 - Consultar videos educativos",
            "UC-08 - Gerenciar configuracoes e preferencias",
        ],
    )

    add_heading(doc, "4. Itens fora de escopo", 1)
    add_bullets(
        doc,
        [
            "Teste de carga e estresse em larga escala",
            "Teste de intrusao de seguranca (pentest)",
            "Testes de compatibilidade com todos os navegadores legacy",
        ],
    )

    add_heading(doc, "5. Ambiente e dados de teste", 1)
    add_bullets(
        doc,
        [
            "Frontend web publicado no Render",
            "API REST Node.js/Express",
            "Banco PostgreSQL no Render",
            "Browsers: Chrome e Firefox",
            "Usuarios de teste com perfis cuidador, familiar e profissional",
            "Dados seed para quiz, guia e videos",
        ],
    )

    add_heading(doc, "6. Criterios de entrada e saida", 1)
    doc.add_paragraph("Criterios de entrada:")
    add_bullets(
        doc,
        [
            "Deploy ativo e endpoint /api/health respondendo status ok",
            "Variaveis de ambiente configuradas (DATABASE_URL, JWT_SECRET, etc.)",
            "Banco acessivel e scripts de inicializacao executados",
        ],
    )
    doc.add_paragraph("Criterios de saida:")
    add_bullets(
        doc,
        [
            "100% dos casos de uso executados ao menos uma vez",
            "Casos criticos sem falha aberta (login, checklist, persistencia)",
            "Evidencias registradas para cada caso (print/log/resultado)",
        ],
    )

    add_heading(doc, "7. Estrategia de teste", 1)
    add_bullets(
        doc,
        [
            "Teste funcional por caso de uso",
            "Teste de integracao frontend + API + banco",
            "Teste de validacao de dados e regras de negocio",
            "Teste de navegacao e experiencia basica do usuario",
            "Teste de regressao rapida apos correcoes",
        ],
    )

    add_heading(doc, "8. Matriz de cobertura por caso de uso", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Caso de uso"
    headers[1].text = "Objetivo"
    headers[2].text = "Tipo"
    headers[3].text = "Prioridade"
    headers[4].text = "Status esperado"

    matrix_rows = [
        ("UC-01", "Criar conta com consentimento", "Funcional/Integracao", "Alta", "Aprovado"),
        ("UC-02", "Autenticar e iniciar sessao", "Funcional/Seguranca", "Alta", "Aprovado"),
        ("UC-03", "Registrar rotina diaria", "Funcional/Integracao", "Alta", "Aprovado"),
        ("UC-04", "Exibir indicadores de adesao", "Funcional", "Media", "Aprovado"),
        ("UC-05", "Exibir etapas do guia", "Funcional", "Media", "Aprovado"),
        ("UC-06", "Corrigir quiz e salvar tentativa", "Funcional/Integracao", "Alta", "Aprovado"),
        ("UC-07", "Listar e reproduzir videos", "Funcional", "Media", "Aprovado"),
        ("UC-08", "Salvar preferencias do usuario", "Funcional/Integracao", "Media", "Aprovado"),
    ]

    for row in matrix_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    add_heading(doc, "9. Casos de teste detalhados", 1)

    add_heading(doc, "9.1 UC-01 - Cadastrar usuario", 2)
    add_case(
        doc,
        "CT-UC01-01",
        "Cadastrar com dados validos",
        "Positivo",
        ["Tela de cadastro aberta"],
        [
            "Informar nome, email valido, perfil e senha valida",
            "Aceitar termos",
            "Enviar cadastro",
        ],
        [
            "Conta criada com sucesso",
            "Token retornado",
            "Usuario persistido no banco",
        ],
    )
    add_case(
        doc,
        "CT-UC01-02",
        "Cadastrar sem aceite dos termos",
        "Negativo",
        ["Tela de cadastro aberta"],
        ["Preencher dados e nao marcar aceite", "Enviar cadastro"],
        ["Cadastro bloqueado", "Mensagem de validacao exibida"],
        "Alta",
    )
    add_case(
        doc,
        "CT-UC01-03",
        "Cadastrar com email ja existente",
        "Negativo",
        ["Ja existe usuario com o email informado"],
        ["Informar email duplicado", "Enviar cadastro"],
        ["Erro de conflito retornado", "Nenhuma duplicidade no banco"],
        "Alta",
    )

    add_heading(doc, "9.2 UC-02 - Realizar login", 2)
    add_case(
        doc,
        "CT-UC02-01",
        "Login com credenciais validas",
        "Positivo",
        ["Usuario previamente cadastrado"],
        ["Informar email e senha validos", "Enviar login"],
        ["Login realizado", "Token e perfil retornados"],
    )
    add_case(
        doc,
        "CT-UC02-02",
        "Login com senha incorreta",
        "Negativo",
        ["Usuario previamente cadastrado"],
        ["Informar email valido e senha incorreta", "Enviar login"],
        ["Acesso negado", "Mensagem de credenciais invalidas"],
    )
    add_case(
        doc,
        "CT-UC02-03",
        "Acesso a rota protegida sem token",
        "Negativo",
        ["Endpoint protegido disponivel"],
        ["Chamar endpoint sem header Authorization"],
        ["Resposta 401", "Dados nao expostos"],
    )

    add_heading(doc, "9.3 UC-03 - Registrar checklist diario", 2)
    add_case(
        doc,
        "CT-UC03-01",
        "Salvar checklist completo",
        "Positivo",
        ["Usuario autenticado"],
        [
            "Selecionar data",
            "Marcar manha, tarde e noite",
            "Selecionar nivel de resistencia",
            "Inserir observacao",
            "Salvar",
        ],
        ["Registro salvo", "Mensagem de sucesso", "Dados persistidos"],
    )
    add_case(
        doc,
        "CT-UC03-02",
        "Atualizar checklist da mesma data",
        "Positivo",
        ["Ja existe checklist na data informada"],
        ["Alterar campos do checklist", "Salvar novamente"],
        ["Registro atualizado", "Sem duplicidade por data"],
    )
    add_case(
        doc,
        "CT-UC03-03",
        "Salvar checklist sem data",
        "Negativo",
        ["Usuario autenticado"],
        ["Deixar data em branco", "Salvar"],
        ["Bloqueio da acao", "Mensagem de data invalida"],
    )

    add_heading(doc, "9.4 UC-04 - Consultar dashboard", 2)
    add_case(
        doc,
        "CT-UC04-01",
        "Carregar indicadores do mes atual",
        "Positivo",
        ["Existem registros de checklist no mes"],
        ["Abrir dashboard"],
        ["Adesao, escovacoes e resistencia exibidos"],
        "Media",
    )
    add_case(
        doc,
        "CT-UC04-02",
        "Trocar mes de referencia",
        "Positivo",
        ["Dashboard aberto"],
        ["Selecionar outro mes"],
        ["Indicadores recalculados para o mes selecionado"],
        "Media",
    )
    add_case(
        doc,
        "CT-UC04-03",
        "Consultar mes sem dados",
        "Negativo",
        ["Nao ha registros no mes selecionado"],
        ["Selecionar mes sem registros"],
        ["Valores zerados sem quebra da interface"],
        "Media",
    )

    add_heading(doc, "9.5 UC-05 - Consultar guia de escovacao", 2)
    add_case(
        doc,
        "CT-UC05-01",
        "Listar etapas do guia",
        "Positivo",
        ["Usuario autenticado"],
        ["Abrir aba Guia"],
        ["Etapas exibidas em ordem", "Titulos, textos e imagens carregados"],
        "Media",
    )
    add_case(
        doc,
        "CT-UC05-02",
        "Visualizacao em mobile",
        "Positivo",
        ["Dispositivo ou viewport mobile"],
        ["Abrir aba Guia em layout mobile"],
        ["Cards legiveis e responsivos"],
        "Baixa",
    )

    add_heading(doc, "9.6 UC-06 - Responder quiz educativo", 2)
    add_case(
        doc,
        "CT-UC06-01",
        "Responder quiz completo",
        "Positivo",
        ["Perguntas carregadas"],
        ["Responder todas as perguntas", "Enviar respostas"],
        ["Pontuacao calculada", "Feedback exibido", "Tentativa salva"],
    )
    add_case(
        doc,
        "CT-UC06-02",
        "Enviar quiz incompleto",
        "Negativo",
        ["Perguntas carregadas"],
        ["Responder apenas parte do quiz", "Enviar"],
        ["Envio bloqueado", "Mensagem solicitando completar respostas"],
        "Alta",
    )
    add_case(
        doc,
        "CT-UC06-03",
        "Consultar historico de tentativas",
        "Positivo",
        ["Ao menos uma tentativa registrada"],
        ["Abrir bloco de historico do quiz"],
        ["Tentativas listadas com data e pontuacao"],
        "Media",
    )

    add_heading(doc, "9.7 UC-07 - Consultar videos educativos", 2)
    add_case(
        doc,
        "CT-UC07-01",
        "Listar videos disponiveis",
        "Positivo",
        ["Usuario autenticado"],
        ["Abrir aba Videos"],
        ["Lista de videos exibida"],
        "Media",
    )
    add_case(
        doc,
        "CT-UC07-02",
        "Reproduzir video incorporado",
        "Positivo",
        ["Lista de videos carregada"],
        ["Iniciar reproducao de um video"],
        ["Player abre e reproduz conteudo"],
        "Baixa",
    )

    add_heading(doc, "9.8 UC-08 - Configuracoes e preferencias", 2)
    add_case(
        doc,
        "CT-UC08-01",
        "Salvar horarios de lembrete",
        "Positivo",
        ["Usuario autenticado"],
        ["Informar horarios HH:MM", "Salvar preferencias"],
        ["Preferencias salvas no banco"],
        "Media",
    )
    add_case(
        doc,
        "CT-UC08-02",
        "Alterar modo de acessibilidade",
        "Positivo",
        ["Usuario autenticado"],
        ["Selecionar alto contraste ou texto ampliado", "Salvar"],
        ["Modo aplicado na interface", "Preferencia persistida"],
        "Media",
    )
    add_case(
        doc,
        "CT-UC08-03",
        "Salvar horario invalido",
        "Negativo",
        ["Usuario autenticado"],
        ["Informar horario fora do formato HH:MM", "Salvar"],
        ["Validacao rejeita formato invalido"],
        "Alta",
    )

    add_heading(doc, "10. Riscos e mitigacoes", 1)
    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = "Table Grid"
    rh = risk_table.rows[0].cells
    rh[0].text = "Risco"
    rh[1].text = "Impacto"
    rh[2].text = "Mitigacao"
    risks = [
        (
            "Indisponibilidade do banco",
            "Falha em login e gravacao",
            "Monitorar conexao, health check e variaveis de ambiente",
        ),
        (
            "Erro de deploy",
            "Sistema fora do ar",
            "Versionamento com tags e rollback controlado",
        ),
        (
            "Quebra de validacoes no frontend",
            "Dados inconsistentes",
            "Executar regressao rapida apos cada alteracao",
        ),
    ]
    for risk in risks:
        cells = risk_table.add_row().cells
        for i, value in enumerate(risk):
            cells[i].text = value

    add_heading(doc, "11. Evidencias de teste", 1)
    add_bullets(
        doc,
        [
            "ID do caso de teste executado",
            "Data e hora da execucao",
            "Resultado (Aprovado/Reprovado)",
            "Print de tela, log ou resposta da API",
            "Observacoes e acao corretiva quando houver falha",
        ],
    )

    add_heading(doc, "12. Cronograma sugerido da execucao", 1)
    add_bullets(
        doc,
        [
            "Dia 1: UC-01, UC-02, UC-03 (fluxos criticos)",
            "Dia 2: UC-04, UC-05, UC-06",
            "Dia 3: UC-07, UC-08, retestes e consolidacao de evidencias",
        ],
    )

    add_heading(doc, "13. Conclusao", 1)
    doc.add_paragraph(
        "Este plano de testes completo amplia a cobertura da entrega anterior e garante "
        "a validacao sistematica de todos os casos de uso do Sorriso Amigo, com foco em "
        "funcionalidade, persistencia de dados, seguranca, experiencia do usuario e "
        "rastreabilidade de resultados para apresentacao academica."
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
